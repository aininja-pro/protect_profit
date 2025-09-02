from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import os
import tempfile
from typing import Dict, Any
from datetime import datetime

class WordRFQGenerator:
    
    def generate_rfq_document(self, scope_data: Dict[str, Any], project_data: Dict[str, Any] = None) -> str:
        """Generate a professional RFQ Word document and return file path"""
        
        # Create document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Header
        header = doc.add_heading('REQUEST FOR QUOTE (RFQ)', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
        
        # Project Information Table
        project_info = doc.add_table(rows=5, cols=2)
        project_info.style = 'Light Grid Accent 1'
        
        # Set column widths
        for row in project_info.rows:
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(4.0)
        
        # Fill project info - handle None values safely
        project_name = '[PROJECT NAME]'
        project_location = '[PROJECT ADDRESS]'
        
        if project_data and isinstance(project_data, dict):
            project_name = project_data.get('name') or '[PROJECT NAME]'
            project_location = project_data.get('location') or '[PROJECT ADDRESS]'
        
        cells = project_info.rows
        cells[0].cells[0].text = 'Project Title:'
        cells[0].cells[1].text = project_name
        cells[1].cells[0].text = 'Project Location:'
        cells[1].cells[1].text = project_location
        cells[2].cells[0].text = 'RFQ Number:'
        scope_id = scope_data.get('id', 'XXXX') if scope_data and isinstance(scope_data, dict) else 'XXXX'
        cells[2].cells[1].text = f"RFQ-{str(scope_id)[:8].upper()}"
        cells[3].cells[0].text = 'Date Issued:'
        cells[3].cells[1].text = datetime.now().strftime('%B %d, %Y')
        cells[4].cells[0].text = 'Submission Deadline:'
        cells[4].cells[1].text = '[SUBMISSION DEADLINE - FILL IN]'
        
        # Make first column bold
        for row in project_info.rows:
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Main content sections - handle None values safely  
        description = ''
        specifications = ''
        exclusions = ''
        
        if scope_data and isinstance(scope_data, dict):
            description = scope_data.get('description') or ''
            specifications = scope_data.get('specifications') or ''
            exclusions = scope_data.get('exclusions') or ''
        
        # Parse and format the AI-enhanced content properly
        if description.strip().startswith('#') or 'DETAILED SCOPE EXPANSION' in description:
            # This is AI-enhanced content, parse it properly
            self._add_ai_enhanced_content(doc, description)
        else:
            # Simple scope content
            self._add_section(doc, "SCOPE OF WORK", description)
        
        if specifications:
            self._add_section(doc, "SPECIFICATIONS", specifications)
        
        if exclusions:
            self._add_section(doc, "EXCLUSIONS", exclusions)
        
        # Standard RFQ sections
        contractor_requirements = """
• Valid contractor's license for the state of [STATE]
• General liability insurance minimum $2,000,000
• Workers' compensation insurance as required by law
• Minimum 5 years experience in similar work
• OSHA compliance and safety plan required
• Required submittals: shop drawings, material data, samples
        """.strip()
        self._add_section(doc, "4. CONTRACTOR REQUIREMENTS", contractor_requirements)
        
        pricing_structure = """
• Provide detailed line-item pricing breakdown
• Include unit prices for materials and labor
• Submit alternate pricing options if applicable
• Payment terms: Net 30 days upon completion of milestones
        """.strip()
        self._add_section(doc, "5. PRICING STRUCTURE", pricing_structure)
        
        timeline = """
• Start Date: [REQUESTED START DATE]
• Completion Date: [REQUESTED COMPLETION DATE]
• Working Hours: Monday-Friday, 7:00 AM to 5:00 PM
• Coordinate with project manager for scheduling
        """.strip()
        self._add_section(doc, "6. TIMELINE & SCHEDULING", timeline)
        
        submission_requirements = """
• Submit proposals in PDF format via email
• Include: detailed proposal, pricing breakdown, references
• All questions must be submitted by [QUESTIONS DEADLINE]
• Quotes must be valid for 60 days minimum
        """.strip()
        self._add_section(doc, "7. SUBMISSION REQUIREMENTS", submission_requirements)
        
        # Contact information
        doc.add_paragraph()
        contact_para = doc.add_paragraph()
        contact_para.add_run('Contact Information:').bold = True
        doc.add_paragraph('[YOUR NAME]\n[YOUR TITLE]\n[YOUR COMPANY]\n[CONTACT PHONE]\n[CONTACT EMAIL]')
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph('We look forward to receiving your detailed proposal and working together on this project. Thank you for your interest and participation.')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.italic = True
        
        # Save to temporary file
        temp_dir = tempfile.gettempdir()
        scope_id_part = str(scope_data.get('id', 'TEMP'))[:8] if scope_data and isinstance(scope_data, dict) else 'TEMP'
        filename = f"RFQ_{scope_id_part}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(temp_dir, filename)
        
        doc.save(file_path)
        return file_path
    
    def _add_section(self, doc: Document, title: str, content: str):
        """Add a formatted section to the document"""
        # Section heading
        heading = doc.add_heading(title, level=2)
        heading.runs[0].font.color.rgb = None  # Keep default color
        
        # Section content
        if not content.strip():
            doc.add_paragraph("[CONTENT TO BE ADDED]")
            doc.add_paragraph()
            return
            
        lines = content.strip().split('\n')
        current_paragraph = ""
        
        for line in lines:
            line = line.strip()
            if not line:
                # Empty line - end current paragraph and start new one
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)
                    current_paragraph = ""
                continue
                
            if line.startswith('•') or line.startswith('-'):
                # Finish any current paragraph
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)
                    current_paragraph = ""
                
                # Add bullet point
                para = doc.add_paragraph()
                para.style = 'List Bullet'
                para.add_run(line.lstrip('•-').strip())
            elif line.endswith(':') or line.isupper():
                # This looks like a subheading
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)
                    current_paragraph = ""
                
                subheading = doc.add_paragraph()
                run = subheading.add_run(line)
                run.bold = True
            else:
                # Regular content line
                if current_paragraph:
                    current_paragraph += " " + line
                else:
                    current_paragraph = line
        
        # Add any remaining paragraph
        if current_paragraph:
            doc.add_paragraph(current_paragraph)
        
        doc.add_paragraph()  # Add spacing

    def _add_ai_enhanced_content(self, doc: Document, content: str):
        """Parse AI-enhanced content and format it properly"""
        lines = content.strip().split('\n')
        current_section = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Skip title lines that duplicate table info
            if any(skip in line.upper() for skip in ['REQUEST FOR QUOTE', 'PROJECT TITLE:', 'PROJECT LOCATION:', 'RFQ NUMBER:', 'DATE ISSUED:', 'SUBMISSION DEADLINE:']):
                continue
                
            # Main section headings (like "### 1. DETAILED SCOPE EXPANSION")
            if line.startswith('###') or (line.startswith('##') and 'DETAILED SCOPE EXPANSION' in line):
                section_title = line.lstrip('#').strip()
                # Remove numbering for cleaner look
                section_title = section_title.replace('1. ', '').replace('2. ', '').replace('3. ', '')
                section_title = section_title.replace('4. ', '').replace('5. ', '').replace('6. ', '')
                section_title = section_title.replace('7. ', '').replace('8. ', '')
                
                heading = doc.add_heading(section_title, level=2)
                current_section = section_title
                continue
                
            # Subsection headings (like "#### 1.1 Site Preparation")  
            elif line.startswith('####') or (line.startswith('**') and line.endswith('**')):
                subsection_title = line.lstrip('#').strip().lstrip('**').rstrip('**')
                # Remove numbering like "1.1" for cleaner look
                import re
                subsection_title = re.sub(r'^\d+\.\d+\s*', '', subsection_title)
                
                subheading = doc.add_paragraph()
                run = subheading.add_run(subsection_title)
                run.bold = True
                subheading.space_after = Pt(6)
                continue
                
            # Bullet points
            elif line.startswith('•') or line.startswith('-'):
                para = doc.add_paragraph()
                para.style = 'List Bullet'
                para.add_run(line.lstrip('•-').strip())
                para.left_indent = Inches(0.25)
                continue
                
            # Regular content paragraphs
            else:
                para = doc.add_paragraph(line)
                if 'Budget Allowance:' in line:
                    para.runs[0].bold = True
        
        doc.add_paragraph()

# Global instance
word_rfq_generator = WordRFQGenerator()